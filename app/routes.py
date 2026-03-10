"""
Flask routes for JobHunter dashboard.
Handles all web interface endpoints and API endpoints for AJAX updates.
"""

import io
import json
import logging
import os
import secrets
import time
import threading
import uuid
from contextlib import contextmanager
from datetime import datetime, timedelta
from pathlib import Path
from urllib.parse import urlparse, urljoin

logger = logging.getLogger(__name__)

# ── Security event logger ──────────────────────────────────────────────────────
_sec_logger = logging.getLogger("jobhunter.security")
_sec_logger.setLevel(logging.INFO)
_sec_logger.propagate = False

def _init_security_logger(log_path: str) -> None:
    """Configure the security logger file handler (called once at app startup)."""
    if _sec_logger.handlers:
        return
    try:
        from pathlib import Path as _P
        _P(log_path).parent.mkdir(parents=True, exist_ok=True)
        fh = logging.FileHandler(log_path, encoding='utf-8')
        fh.setFormatter(logging.Formatter('%(message)s'))
        _sec_logger.addHandler(fh)
    except OSError as e:
        logger.warning("Cannot open security log %s: %s", log_path, e)


def _sec_log(event_type: str, username: str = "-", details: str = "-") -> None:
    """Write one security event line: timestamp | event_type | ip | username | details."""
    from flask import request as _req
    # ProxyFix is configured in create_app() — remote_addr already reflects the
    # real client IP (set by Nginx via X-Forwarded-For, trusted exactly 1 hop).
    ip = _req.remote_addr or "-"
    ts = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    _sec_logger.info("%s | %-22s | %-15s | %-20s | %s", ts, event_type, ip, username, details)

# Cross-process file locking (Linux/Mac only; Windows falls back to thread lock)
try:
    import fcntl as _fcntl
    _HAS_FCNTL = True
except ImportError:
    _HAS_FCNTL = False

from flask import Blueprint, render_template, request, jsonify, session, redirect, url_for, send_file
from sqlalchemy import func, text
from sqlalchemy.orm import joinedload

from werkzeug.utils import secure_filename

from app.database import SessionLocal
from app.models import Offer, Tracking, Domain, User, UserOffer, PasswordReset, EmailConfirmation
from app.auth import login_required, admin_required, superadmin_required, check_credentials, get_current_role
from app.services.filter_engine import normalize_text
from config import TARGET_COMPANIES, DATA_DIR
from app import limiter

# Create Blueprint
bp = Blueprint('main', __name__)

# Track application start time for /health uptime
_APP_START_TIME = datetime.utcnow()

# ── Password policy ───────────────────────────────────────────────────────────

def _validate_password(password: str) -> list[str]:
    """Return a list of policy violation messages (empty = OK).
    Policy: 12+ chars, at least 1 uppercase letter, at least 1 digit.
    """
    errs = []
    if len(password) < 12:
        errs.append("Mot de passe trop court (12 caractères minimum).")
    if len(password) > 72:
        errs.append("Mot de passe trop long (72 caractères maximum).")
    if not any(c.isupper() for c in password):
        errs.append("Le mot de passe doit contenir au moins une lettre majuscule.")
    if not any(c.isdigit() for c in password):
        errs.append("Le mot de passe doit contenir au moins un chiffre.")
    return errs


# ── TOTP secret encryption (Fernet) ───────────────────────────────────────────

def _get_fernet():
    """Return a Fernet instance if TOTP_ENCRYPTION_KEY is configured, else None."""
    from config import Config
    key = Config.TOTP_ENCRYPTION_KEY
    if not key:
        return None
    try:
        from cryptography.fernet import Fernet
        return Fernet(key.encode() if isinstance(key, str) else key)
    except Exception:
        return None


def _encrypt_totp_secret(plaintext: str) -> str:
    """Encrypt a TOTP secret. Returns encrypted token string, or plaintext if no key."""
    f = _get_fernet()
    if f is None:
        return plaintext
    return f.encrypt(plaintext.encode()).decode()


def _decrypt_totp_secret(value: str) -> str:
    """Decrypt a TOTP secret. Handles both encrypted and legacy plaintext values."""
    if not value:
        return value
    f = _get_fernet()
    if f is None:
        return value
    try:
        from cryptography.fernet import InvalidToken
        return f.decrypt(value.encode()).decode()
    except (InvalidToken, Exception):
        # Plaintext legacy secret (not yet migrated)
        return value

# Predefined security questions for account recovery
SECURITY_QUESTIONS = [
    "Quel était le nom de votre premier animal de compagnie ?",
    "Quelle est votre ville de naissance ?",
    "Quel était le nom de votre meilleur ami d'enfance ?",
    "Quel était le nom de votre école primaire ?",
    "Quel est votre plat préféré ?",
]

# ── User activity helpers ──────────────────────────────────────────────────────

def _touch_last_login(user_id: int) -> None:
    """Update last_login timestamp for a DB user.

    Uses SessionLocal.session_factory() to create an isolated (non-scoped) session
    so it never closes the thread-local scoped session that a calling route may hold.
    """
    try:
        db = SessionLocal.session_factory()
        try:
            db.query(User).filter(User.id == user_id).update(
                {"last_login": datetime.utcnow()},
                synchronize_session=False,
            )
            db.commit()
        finally:
            db.close()
    except Exception as exc:
        logger.warning("_touch_last_login failed: %s", exc)


def _add_claude_tokens(user_id: int, tokens: int) -> None:
    """Atomically add *tokens* to user.claude_tokens_used.

    Uses SessionLocal.session_factory() to create an isolated (non-scoped) session
    so it never closes the thread-local scoped session that a calling route may hold.
    """
    if not user_id or tokens <= 0:
        return
    try:
        db = SessionLocal.session_factory()
        try:
            db.execute(
                text("UPDATE users SET claude_tokens_used = claude_tokens_used + :t WHERE id = :uid"),
                {"t": tokens, "uid": user_id},
            )
            db.commit()
        finally:
            db.close()
    except Exception as exc:
        logger.warning("_add_claude_tokens failed: %s", exc)


# ── Weekly quota helpers ───────────────────────────────────────────────────────

WEEKLY_MATCH_LIMIT  = 1   # max AI matchings per week for role=user
WEEKLY_LETTER_LIMIT = 5   # max cover letter generations per week for role=user


def _next_monday() -> datetime:
    """Return next Monday at 00:00 UTC."""
    now = datetime.utcnow()
    days_ahead = (7 - now.weekday()) % 7  # 0 = already Monday
    if days_ahead == 0:
        days_ahead = 7
    return (now + timedelta(days=days_ahead)).replace(hour=0, minute=0, second=0, microsecond=0)


def _check_and_increment_quota(user_id: int, quota_field: str, limit: int) -> tuple[bool, int, int]:
    """Check and increment a weekly quota for a DB user (role=user only).

    Returns (allowed, used_after, limit).
    Admin users (role != 'user') are always allowed.
    """
    if not user_id:
        return True, 0, limit  # config/legacy admin — no limit

    try:
        db = SessionLocal.session_factory()
        try:
            user = db.query(User).filter(User.id == user_id).first()
            if not user:
                return True, 0, limit
            if user.role != 'user':
                return True, 0, limit  # admin / viewer — no quota

            now = datetime.utcnow()
            # Reset counters if quota_reset_at is in the past (or unset)
            if not user.quota_reset_at or user.quota_reset_at <= now:
                user.weekly_matches_used = 0
                user.weekly_letters_used = 0
                user.quota_reset_at = _next_monday()

            current = getattr(user, quota_field, 0) or 0
            if current >= limit:
                db.commit()
                return False, current, limit

            setattr(user, quota_field, current + 1)
            db.commit()
            return True, current + 1, limit
        finally:
            db.close()
    except Exception as exc:
        logger.warning("_check_and_increment_quota failed: %s", exc)
        return True, 0, limit  # fail open (don't block user on DB error)


def _get_user_quota(user_id: int) -> dict:
    """Return current quota state for a DB user.  Returns None for admin/config user."""
    if not user_id:
        return None
    try:
        db = SessionLocal.session_factory()
        try:
            user = db.query(User).filter(User.id == user_id).first()
            if not user or user.role != 'user':
                return None
            now = datetime.utcnow()
            # If reset_at is past, quotas are effectively 0
            if not user.quota_reset_at or user.quota_reset_at <= now:
                return {
                    'matches_used': 0, 'matches_limit': WEEKLY_MATCH_LIMIT,
                    'letters_used': 0, 'letters_limit': WEEKLY_LETTER_LIMIT,
                }
            return {
                'matches_used': user.weekly_matches_used or 0,
                'matches_limit': WEEKLY_MATCH_LIMIT,
                'letters_used': user.weekly_letters_used or 0,
                'letters_limit': WEEKLY_LETTER_LIMIT,
            }
        finally:
            db.close()
    except Exception:
        return None


# ── Async CV matching task registry (file-backed, multi-worker safe) ──────────
# Tasks are stored in data/matching_tasks.json so all Gunicorn workers share state.
# File is keyed by str(user_id) or "_admin" for config/legacy admin (user_id=None).
_task_thread_lock = threading.Lock()  # intra-process thread safety


_REDIRECT_WHITELIST = (
    "/dashboard",
    "/stats",
    "/documents",
    "/account",
    "/account/profile",
    "/admin",
    "/offers",
)


def _is_safe_redirect(url: str) -> bool:
    """Return True only if *url* is a whitelisted relative path on the same host."""
    if not url:
        return False
    ref = urlparse(request.host_url)
    test = urlparse(urljoin(request.host_url, url))
    if not (test.scheme in ("http", "https") and ref.netloc == test.netloc):
        return False
    path = test.path.rstrip("/") or "/"
    return any(path == w or path.startswith(w + "/") for w in _REDIRECT_WHITELIST)


# ── Auth routes ───────────────────────────────────────────────────────────────

@bp.route('/login', methods=['GET', 'POST'])
@limiter.limit("20 per minute")
def login():
    """Login page. Redirects to dashboard if already authenticated."""
    if session.get("username"):
        return redirect(url_for("main.dashboard"))

    error = None
    confirm_resend_user_id = None
    if request.method == 'POST':
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        role, user_id, domain_id = check_credentials(username, password)
        if role == "inactive":
            # user_id is set, domain_id is actually email_confirmed (bool)
            email_confirmed = domain_id
            _sec_log("LOGIN_FAIL_INACTIVE", username)
            if not email_confirmed:
                error = "Veuillez d'abord confirmer votre adresse email."
                confirm_resend_user_id = user_id
            else:
                error = "Votre compte a été désactivé. Contactez un administrateur."
        elif role:
            _sec_log("LOGIN_SUCCESS", username, f"role={role}")
            next_url = request.args.get("next", "")
            if not _is_safe_redirect(next_url):
                next_url = url_for("main.dashboard")
            # Check if 2FA is required for this DB user
            if user_id is not None:
                db = SessionLocal()
                try:
                    _u = db.query(User).filter(User.id == user_id).first()
                    if _u and _u.totp_enabled:
                        session.clear()
                        session["_2fa_uid"] = user_id
                        session["_2fa_next"] = next_url
                        return redirect(url_for("main.login_2fa"))
                finally:
                    db.close()
            # No 2FA — complete login immediately
            if user_id is not None:
                _touch_last_login(user_id)
            session.clear()
            session["username"] = username
            session["role"] = role
            session["user_id"] = user_id
            session["domain_id"] = domain_id
            return redirect(next_url)
        else:
            _sec_log("LOGIN_FAIL", username or "-")
            error = "Identifiant ou mot de passe incorrect."

    return render_template("login.html", error=error,
                           confirm_resend_user_id=confirm_resend_user_id)


@bp.route('/login/2fa', methods=['GET', 'POST'])
@limiter.limit("10 per minute")
def login_2fa():
    """Second step of login: verify TOTP code."""
    uid = session.get("_2fa_uid")
    if not uid:
        return redirect(url_for("main.login"))

    error = None
    if request.method == 'POST':
        import pyotp
        code = request.form.get("totp_code", "").strip().replace(" ", "")
        db = SessionLocal()
        try:
            user = db.query(User).filter(User.id == uid).first()
            if user and user.totp_enabled and user.totp_secret:
                totp = pyotp.TOTP(_decrypt_totp_secret(user.totp_secret))
                if totp.verify(code, valid_window=0):
                    next_url = session.pop("_2fa_next", None) or url_for("main.dashboard")
                    uid_val       = user.id
                    username_val  = user.username
                    role_val      = user.role
                    domain_id_val = user.domain_id
                    _sec_log("2FA_SUCCESS", username_val)
                    _touch_last_login(uid_val)
                    session.clear()
                    session["username"]  = username_val
                    session["role"]      = role_val
                    session["user_id"]   = uid_val
                    session["domain_id"] = domain_id_val
                    return redirect(next_url)
                _sec_log("2FA_FAIL", user.username)
                error = "Code A2F invalide. Vérifiez votre application et réessayez."
            else:
                return redirect(url_for("main.login"))
        finally:
            db.close()

    return render_template("login_2fa.html", error=error)


@bp.route('/logout', methods=['GET', 'POST'])
def logout():
    """Clear session and redirect to login."""
    session.clear()
    return redirect(url_for("main.login"))


@bp.route('/register', methods=['GET', 'POST'])
@limiter.limit("10 per minute")
def register():
    """Self-registration: create a new user account linked to a domain."""
    if session.get("username"):
        return redirect(url_for("main.dashboard"))

    db = SessionLocal()
    try:
        domains = db.query(Domain).order_by(Domain.name).all()
        errors = []

        if request.method == 'POST':
            import re as _re
            username = request.form.get("username", "").strip()
            email = request.form.get("email", "").strip().lower()
            password = request.form.get("password", "")
            confirm = request.form.get("confirm_password", "")
            domain_id_raw = request.form.get("domain_id", "").strip()
            security_question = request.form.get("security_question", "").strip()
            security_answer = request.form.get("security_answer", "").strip()

            if not username:
                errors.append("Nom d'utilisateur requis.")
            elif len(username) > 64:
                errors.append("Nom d'utilisateur trop long (max 64 caractères).")
            elif not all(c.isalnum() or c in "-_." for c in username):
                errors.append("L'identifiant ne peut contenir que des lettres, chiffres, tirets, points et underscores.")
            if not email:
                errors.append("Adresse email requise.")
            elif not _re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email):
                errors.append("Adresse email invalide.")
            elif len(email) > 255:
                errors.append("Adresse email trop longue.")
            if not password:
                errors.append("Mot de passe requis.")
            else:
                errors.extend(_validate_password(password))
            if password and password != confirm:
                errors.append("Les mots de passe ne correspondent pas.")
            if not domain_id_raw:
                errors.append("Veuillez choisir un domaine.")
            else:
                try:
                    domain_id_raw = int(domain_id_raw)
                except ValueError:
                    errors.append("Domaine invalide.")
                    domain_id_raw = None
            if not security_question or security_question not in SECURITY_QUESTIONS:
                errors.append("Veuillez choisir une question de sécurité.")
            if not security_answer:
                errors.append("La réponse à la question de sécurité est requise.")
            if not request.form.get("accept_terms"):
                errors.append("Vous devez accepter les CGU et la Politique de confidentialité pour créer un compte.")

            if not errors:
                existing = db.query(User).filter(User.username == username).first()
                if existing:
                    errors.append("Ce nom d'utilisateur est déjà pris.")
                else:
                    existing_email = db.query(User).filter(User.email == email).first()
                    if existing_email:
                        errors.append("Cette adresse email est déjà utilisée.")

            if not errors:
                    from app import bcrypt
                    pw_hash = bcrypt.generate_password_hash(password).decode("utf-8")
                    answer_hash = bcrypt.generate_password_hash(security_answer.lower()).decode("utf-8")
                    new_user = User(
                        username=username,
                        password_hash=pw_hash,
                        role="user",
                        domain_id=domain_id_raw,
                        email=email,
                        is_active=False,
                        email_confirmed=False,
                        security_question=security_question,
                        security_answer_hash=answer_hash,
                    )
                    db.add(new_user)
                    db.flush()
                    # Generate email confirmation token
                    token = secrets.token_urlsafe(32)
                    confirmation = EmailConfirmation(
                        user_id=new_user.id,
                        token=token,
                    )
                    db.add(confirmation)
                    db.commit()
                    # Send confirmation email
                    confirm_url = url_for("main.confirm_email", token=token,
                                          _external=True, _scheme='https')
                    _send_confirmation_email(email, username, confirm_url)
                    _sec_log("REGISTER", username, f"email={email}")
                    return redirect(url_for("main.register_pending", uid=new_user.id))

        return render_template("register.html", domains=domains, errors=errors,
                               security_questions=SECURITY_QUESTIONS)
    finally:
        db.close()


@bp.route('/register/pending')
@limiter.limit("20 per minute")
def register_pending():
    """Post-registration page: email confirmation pending."""
    uid = request.args.get("uid", type=int)
    return render_template("register_pending.html", uid=uid)


@bp.route('/confirm-email/<token>')
@limiter.limit("10 per minute")
def confirm_email(token):
    """Confirm a user's email address via the token sent at registration."""
    db = SessionLocal()
    try:
        confirmation = db.query(EmailConfirmation).filter(
            EmailConfirmation.token == token,
            EmailConfirmation.used == False,
        ).first()
        if not confirmation:
            return render_template("confirm_email.html",
                                   error="Lien invalide ou déjà utilisé.")
        # Check 24h expiry
        from datetime import timezone as _tz
        age = datetime.now(_tz.utc) - confirmation.created_at.replace(tzinfo=_tz.utc)
        if age > timedelta(hours=24):
            confirmation.used = True
            db.commit()
            return render_template("confirm_email.html",
                                   error="Ce lien a expiré (validité 24 heures). Veuillez vous réinscrire.")
        user = db.query(User).filter(User.id == confirmation.user_id).first()
        if not user:
            return render_template("confirm_email.html",
                                   error="Compte introuvable.")
        user.is_active = True
        user.email_confirmed = True
        confirmation.used = True
        db.commit()
        _sec_log("EMAIL_CONFIRMED", user.username)
        return redirect(url_for("main.login") + "?ok=email_confirmed")
    finally:
        db.close()


@bp.route('/api/resend-confirmation', methods=['POST'])
@limiter.limit("1 per minute")
def resend_confirmation():
    """Resend email confirmation (rate limited to 1/min)."""
    if request.is_json:
        uid = request.json.get("user_id")
    else:
        uid = request.form.get("user_id", type=int)
    if not uid:
        return jsonify({"error": "Paramètre manquant"}), 400
    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == uid).first()
        if not user or user.email_confirmed or user.is_active:
            # Don't reveal whether user exists
            return jsonify({"ok": True, "message": "Si le compte existe, un email a été envoyé."})
        if not user.email:
            return jsonify({"error": "Aucune adresse email associée à ce compte."}), 400
        # Invalidate old tokens
        db.query(EmailConfirmation).filter(
            EmailConfirmation.user_id == uid,
            EmailConfirmation.used == False,
        ).update({"used": True})
        # Create new token
        token = secrets.token_urlsafe(32)
        confirmation = EmailConfirmation(user_id=uid, token=token)
        db.add(confirmation)
        db.commit()
        confirm_url = url_for("main.confirm_email", token=token,
                              _external=True, _scheme='https')
        _send_confirmation_email(user.email, user.username, confirm_url)
        return jsonify({"ok": True, "message": "Email de confirmation renvoyé."})
    finally:
        db.close()


VALID_STATUSES = [
    'New', 'Applied', 'Followed up', 'Interview',
    'Accepted', 'Rejected', 'No response',
]

# CV storage paths
CV_DIR = DATA_DIR / "cv"
CV_TEXT_PATH = CV_DIR / "cv_text.txt"

# Task registry file (shared across all Gunicorn workers)
_TASKS_FILE = DATA_DIR / "matching_tasks.json"
_TASKS_LOCK_FILE = DATA_DIR / "matching_tasks.json.lock"


@contextmanager
def _tasks_lock_ctx():
    """Acquire thread lock + optional cross-process file lock (fcntl on Linux)."""
    with _task_thread_lock:
        if not _HAS_FCNTL:
            yield
            return
        try:
            _TASKS_LOCK_FILE.parent.mkdir(parents=True, exist_ok=True)
            lock_fh = open(_TASKS_LOCK_FILE, 'a')
            _fcntl.flock(lock_fh.fileno(), _fcntl.LOCK_EX)
            try:
                yield
            finally:
                _fcntl.flock(lock_fh.fileno(), _fcntl.LOCK_UN)
                lock_fh.close()
        except OSError:
            yield  # fallback: no file lock


def _read_tasks_raw() -> dict:
    """Read the tasks JSON file. Returns {} if missing or unreadable."""
    try:
        if _TASKS_FILE.exists():
            return json.loads(_TASKS_FILE.read_text(encoding='utf-8'))
    except (json.JSONDecodeError, OSError):
        pass
    return {}


def _write_tasks_raw(data: dict) -> None:
    """Atomically write task data via a temp file, cleaning up on failure."""
    import tempfile
    _TASKS_FILE.parent.mkdir(parents=True, exist_ok=True)
    fd, tmp_path = tempfile.mkstemp(dir=str(_TASKS_FILE.parent), suffix='.tmp')
    try:
        with os.fdopen(fd, 'w', encoding='utf-8') as fh:
            fh.write(json.dumps(data, default=str))
        os.replace(tmp_path, str(_TASKS_FILE))
    except Exception:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
        raise


def _load_task(user_id) -> dict | None:
    """Return the task state for user_id, or None if no task exists."""
    key = str(user_id) if user_id is not None else '_admin'
    with _tasks_lock_ctx():
        return _read_tasks_raw().get(key)


def _save_task(user_id, state: dict) -> None:
    """Persist task state for user_id to the shared JSON file."""
    key = str(user_id) if user_id is not None else '_admin'
    with _tasks_lock_ctx():
        data = _read_tasks_raw()
        data[key] = state
        _write_tasks_raw(data)


def _try_start_task(user_id, initial_state: dict) -> bool:
    """
    Atomic check-and-set: if no task is currently running for user_id,
    write initial_state and return True. Returns False if already running.
    """
    key = str(user_id) if user_id is not None else '_admin'
    with _tasks_lock_ctx():
        data = _read_tasks_raw()
        if data.get(key, {}).get('status') == 'running':
            return False
        data[key] = initial_state
        _write_tasks_raw(data)
    return True

# ── Document upload validation ────────────────────────────────────────────────
MAX_UPLOAD_SIZE = 5 * 1024 * 1024  # 5 MB

# Exact MIME types accepted per extension (whitelist)
_ALLOWED_MIMES: dict[str, set[str]] = {
    '.pdf':  {'application/pdf', 'application/octet-stream'},
    '.docx': {
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/zip',
        'application/octet-stream',
    },
    '.txt':  {'text/plain', 'application/octet-stream'},
}

# Always-rejected MIME types (execution risk / active content)
_BLOCKED_MIMES = {
    'text/html', 'application/xhtml+xml',
    'application/x-msdownload', 'application/x-executable',
    'application/x-sh', 'application/x-bat',
}

# Magic bytes signatures for allowed file types
_MAGIC_BYTES: dict[str, bytes] = {
    '.pdf':  b'%PDF',
    '.docx': b'PK\x03\x04',  # ZIP container (DOCX/XLSX/PPTX)
}


def _check_magic_bytes(raw: bytes, ext: str) -> bool:
    """Return True if the file's magic bytes match the declared extension.
    TXT has no magic bytes — we verify it's valid UTF-8 instead.
    """
    if ext in _MAGIC_BYTES:
        return raw[:len(_MAGIC_BYTES[ext])] == _MAGIC_BYTES[ext]
    if ext == '.txt':
        try:
            raw.decode('utf-8')
            return True
        except UnicodeDecodeError:
            return False
    return True  # Unknown extension — skip check


def _user_docs_dir():
    """Return the document directory scoped to the current user session."""
    user_id = session.get("user_id")
    folder = str(user_id) if user_id is not None else "legacy_admin"
    return DATA_DIR / "documents" / folder


_CV_ALLOWED_EXTS = {'.pdf', '.docx', '.txt'}


def _extract_doc_text(path: Path) -> str | None:
    """Extract plain text from a PDF, DOCX, or TXT file.  Returns None on failure."""
    ext = path.suffix.lower()
    try:
        if ext == '.pdf':
            import PyPDF2
            reader = PyPDF2.PdfReader(str(path))
            text = "\n".join(p.extract_text() or "" for p in reader.pages)
            return text.strip() or None
        elif ext == '.docx':
            from docx import Document as _DocxDoc
            doc = _DocxDoc(str(path))
            text = "\n".join(p.text for p in doc.paragraphs)
            return text.strip() or None
        elif ext == '.txt':
            try:
                return path.read_text(encoding='utf-8').strip() or None
            except UnicodeDecodeError:
                return path.read_text(encoding='latin-1', errors='replace').strip() or None
    except Exception as exc:
        logger.warning("_extract_doc_text failed for %s: %s", path, exc)
    return None


def _find_cv_text(user_id) -> str | None:
    """Find and extract CV text for a user.

    Search order for DB users (user_id is not None):
      1. Most recent file whose stem contains 'cv' (case-insensitive) in data/documents/{user_id}/
      2. Most recent PDF or DOCX in that folder
    Falls back to the legacy CV_TEXT_PATH for config/legacy admin (user_id is None).
    Returns None if nothing is found or text extraction fails.
    """
    if user_id is None:
        # Legacy config admin: use the old pre-extracted text file
        if CV_TEXT_PATH.exists():
            return CV_TEXT_PATH.read_text(encoding='utf-8')
        return None

    docs_dir = DATA_DIR / "documents" / str(user_id)
    if not docs_dir.exists():
        return None

    files = [f for f in docs_dir.iterdir() if f.is_file() and f.suffix.lower() in _CV_ALLOWED_EXTS]
    if not files:
        return None

    # Priority 1: files whose stem contains 'cv' (case-insensitive), most recent first
    cv_named = sorted(
        [f for f in files if 'cv' in f.stem.lower()],
        key=lambda f: f.stat().st_mtime,
        reverse=True,
    )
    if cv_named:
        text = _extract_doc_text(cv_named[0])
        if text:
            return text

    # Priority 2: most recent PDF or DOCX
    doc_files = sorted(
        [f for f in files if f.suffix.lower() in ('.pdf', '.docx')],
        key=lambda f: f.stat().st_mtime,
        reverse=True,
    )
    if doc_files:
        return _extract_doc_text(doc_files[0])

    return None


def _has_cv_file(user_id) -> bool:
    """Return True if the user has at least one usable document for CV matching."""
    if user_id is None:
        return CV_TEXT_PATH.exists()

    docs_dir = DATA_DIR / "documents" / str(user_id)
    if not docs_dir.exists():
        return False
    return any(
        f.is_file() and f.suffix.lower() in _CV_ALLOWED_EXTS
        for f in docs_dir.iterdir()
    )


@bp.route('/')
def landing():
    """Public landing page. Redirects to dashboard if already authenticated."""
    if session.get("username"):
        return redirect(url_for("main.dashboard"))
    db = SessionLocal()
    try:
        total_offers = db.query(func.count(Offer.id)).scalar() or 0
        sources_count = db.query(func.count(func.distinct(Offer.source))).scalar() or 0
        domains = db.query(Domain).order_by(Domain.name).all()
        domains_count = len(domains)
        return render_template(
            'landing.html',
            total_offers=total_offers,
            sources_count=sources_count,
            domains_count=domains_count,
            domains=domains,
        )
    finally:
        db.close()


@bp.route('/dashboard')
@login_required
def dashboard():
    """
    Main dashboard view.
    Displays all job offers with their tracking status in an interactive table.
    """
    db = SessionLocal()
    user_id = session.get("user_id")
    domain_id = session.get("domain_id")
    try:
        # Filter offers by domain if the user has one
        query = db.query(Offer)
        if domain_id:
            query = query.filter(Offer.domain_id == domain_id)

        # Build user_offers_map: offer_id -> tracking object
        if user_id is None:
            # Config/legacy user (admin): load via Offer.tracking relationship
            offers = query.options(joinedload(Offer.tracking)).all()
            user_offers_map = {o.id: o.tracking for o in offers if o.tracking}
        else:
            # DB user: load UserOffer rows for this user
            offers = query.all()
            offer_ids = [o.id for o in offers]
            if offer_ids:
                user_offer_rows = db.query(UserOffer).filter(
                    UserOffer.user_id == user_id,
                    UserOffer.offer_id.in_(offer_ids),
                ).all()
            else:
                user_offer_rows = []
            user_offers_map = {uo.offer_id: uo for uo in user_offer_rows}

        total_offers = len(offers)

        # Stats — same fields regardless of tracking backend
        uo_values = list(user_offers_map.values())
        cv_sent_count = sum(1 for uo in uo_values if uo.cv_sent)
        follow_up_count = sum(1 for uo in uo_values if uo.follow_up_done)
        interview_count = sum(1 for uo in uo_values if uo.status == 'Interview')

        stats = {
            'total_offers': total_offers,
            'cv_sent': cv_sent_count,
            'follow_ups': follow_up_count,
            'interviews': interview_count,
        }

        # Collect unique sources for filter dropdown
        sources = sorted(set(o.source for o in offers))

        # Mark target company offers
        targets_norm = [normalize_text(c) for c in TARGET_COMPANIES]
        target_ids = set()
        for o in offers:
            co = normalize_text(o.company or "")
            for t in targets_norm:
                if t in co:
                    target_ids.add(o.id)
                    break

        has_cv = _has_cv_file(user_id)
        cutoff_new = datetime.utcnow() - timedelta(hours=24)

        # Pass domain list to template only for admins (no domain scoping)
        admin_domains = []
        if not domain_id:
            admin_domains = db.query(Domain).order_by(Domain.name).all()

        user_quota = _get_user_quota(user_id)

        return render_template(
            'dashboard.html',
            offers=offers,
            user_offers_map=user_offers_map,
            stats=stats,
            sources=sources,
            statuses=VALID_STATUSES,
            target_ids=target_ids,
            has_cv=has_cv,
            cutoff_new=cutoff_new,
            role=get_current_role(),
            username=session.get("username"),
            admin_domains=admin_domains,
            user_quota=user_quota,
        )
    finally:
        db.close()


@bp.route('/api/tracking/<int:offer_id>', methods=['PUT'])
@login_required
def update_tracking(offer_id):
    """
    AJAX endpoint to update tracking data for an offer.
    Accepts JSON with any combination of: status, cv_sent, follow_up_done,
    date_sent, follow_up_date, notes.
    DB users use UserOffer; config/legacy admin uses Tracking.
    """
    if session.get("role") == "viewer":
        return jsonify({"error": "Accès réservé"}), 403

    t_start = time.perf_counter()
    db = SessionLocal()
    user_id = session.get("user_id")
    try:
        t_q0 = time.perf_counter()
        if user_id is not None:
            tracking = db.query(UserOffer).filter(
                UserOffer.user_id == user_id,
                UserOffer.offer_id == offer_id,
            ).first()
            if not tracking:
                offer_exists = db.query(Offer.id).filter(Offer.id == offer_id).scalar()
                if not offer_exists:
                    return jsonify({'error': 'Offer not found'}), 404
                tracking = UserOffer(user_id=user_id, offer_id=offer_id, status='New')
                db.add(tracking)
        else:
            tracking = db.query(Tracking).filter(Tracking.offer_id == offer_id).first()
            if not tracking:
                offer_exists = db.query(Offer.id).filter(Offer.id == offer_id).scalar()
                if not offer_exists:
                    return jsonify({'error': 'Offer not found'}), 404
                tracking = Tracking(offer_id=offer_id, status='New')
                db.add(tracking)
        t_q1 = time.perf_counter()

        data = request.get_json()

        t_upd0 = time.perf_counter()
        if 'status' in data:
            if data['status'] in VALID_STATUSES:
                tracking.status = data['status']

        if 'cv_sent' in data:
            tracking.cv_sent = bool(data['cv_sent'])
            if tracking.cv_sent and not tracking.date_sent:
                tracking.date_sent = datetime.utcnow()
            elif not tracking.cv_sent:
                tracking.date_sent = None

        if 'follow_up_done' in data:
            tracking.follow_up_done = bool(data['follow_up_done'])
            if tracking.follow_up_done and not tracking.follow_up_date:
                tracking.follow_up_date = datetime.utcnow()
            elif not tracking.follow_up_done:
                tracking.follow_up_date = None

        if 'notes' in data:
            tracking.notes = data['notes'].strip() if data['notes'] else None

        tracking.updated_at = datetime.utcnow()
        t_upd1 = time.perf_counter()

        t_c0 = time.perf_counter()
        db.commit()
        t_c1 = time.perf_counter()

        t_total = (time.perf_counter() - t_start) * 1000

        return jsonify({
            'ok': True,
            'server_ms': round(t_total, 1),
            'tracking': {
                'status': tracking.status,
                'cv_sent': tracking.cv_sent,
                'follow_up_done': tracking.follow_up_done,
                'date_sent': tracking.date_sent.strftime('%Y-%m-%d') if tracking.date_sent else None,
                'follow_up_date': tracking.follow_up_date.strftime('%Y-%m-%d') if tracking.follow_up_date else None,
                'notes': tracking.notes,
            }
        })

    except Exception:
        db.rollback()
        return jsonify({'error': 'Erreur interne du serveur'}), 500
    finally:
        db.close()


@bp.route('/api/tracking/<int:offer_id>/favorite', methods=['POST'])
@login_required
def toggle_favorite(offer_id):
    """Toggle the is_favorite flag for a user's offer."""
    user_id = session.get("user_id")
    if user_id is None:
        return jsonify({'error': 'Non disponible'}), 400
    if session.get("role") == "viewer":
        return jsonify({"error": "Accès réservé"}), 403

    db = SessionLocal()
    try:
        uo = db.query(UserOffer).filter(
            UserOffer.user_id == user_id,
            UserOffer.offer_id == offer_id,
        ).first()
        if not uo:
            offer_exists = db.query(Offer.id).filter(Offer.id == offer_id).scalar()
            if not offer_exists:
                return jsonify({'error': 'Offer not found'}), 404
            uo = UserOffer(user_id=user_id, offer_id=offer_id, status='New', is_favorite=True)
            db.add(uo)
        else:
            uo.is_favorite = not uo.is_favorite
        uo.updated_at = datetime.utcnow()
        db.commit()
        return jsonify({'ok': True, 'is_favorite': uo.is_favorite})
    except Exception:
        db.rollback()
        return jsonify({'error': 'Erreur interne du serveur'}), 500
    finally:
        db.close()


@bp.route('/offer/<int:offer_id>')
@login_required
def offer_detail(offer_id):
    """Detailed view of a single job offer."""
    db = SessionLocal()
    user_id = session.get("user_id")
    domain_id = session.get("domain_id")
    try:
        offer = db.query(Offer).options(joinedload(Offer.tracking)).filter(
            Offer.id == offer_id
        ).first()
        if not offer:
            return "Offer not found", 404

        # Domain authorization: domain-scoped users may only view offers in their domain
        if domain_id and offer.domain_id and offer.domain_id != domain_id:
            return "Accès refusé", 403

        # Resolve per-user tracking object
        if user_id is not None:
            user_offer = db.query(UserOffer).filter(
                UserOffer.user_id == user_id,
                UserOffer.offer_id == offer_id,
            ).first()
        else:
            user_offer = offer.tracking  # config admin uses Tracking

        docs_dir = _user_docs_dir()
        docs_dir.mkdir(parents=True, exist_ok=True)
        doc_files = sorted(f.name for f in docs_dir.iterdir() if f.is_file())
        return render_template('offer_detail.html', offer=offer,
                               user_offer=user_offer,
                               doc_files=doc_files,
                               role=get_current_role(),
                               username=session.get("username"))
    finally:
        db.close()


@bp.route('/stats')
@login_required
def stats():
    """Statistics page with detailed metrics."""
    db = SessionLocal()
    user_id = session.get("user_id")
    domain_id = session.get("domain_id")
    try:
        offer_query = db.query(Offer)
        if domain_id:
            offer_query = offer_query.filter(Offer.domain_id == domain_id)
        total_offers = offer_query.count()

        def _uo():
            """Base query for the current user's tracking rows."""
            if user_id is not None:
                return db.query(UserOffer).filter(UserOffer.user_id == user_id)
            return db.query(Tracking)

        tracked_offers = _uo().count()
        cv_sent = _uo().filter(
            (UserOffer.cv_sent if user_id is not None else Tracking.cv_sent) == True
        ).count()
        follow_ups = _uo().filter(
            (UserOffer.follow_up_done if user_id is not None else Tracking.follow_up_done) == True
        ).count()

        status_counts = {}
        status_col = UserOffer.status if user_id is not None else Tracking.status
        for status in VALID_STATUSES:
            count = _uo().filter(status_col == status).count()
            status_counts[status] = count

        # Interviews count for response rate
        interviews = status_counts.get('Interview', 0) + status_counts.get('Accepted', 0)
        response_rate = round(interviews / cv_sent * 100, 1) if cv_sent > 0 else 0

        # Average CV match score
        avg_cv_score = 0.0
        high_score_count = 0
        if user_id is not None:
            avg_row = db.query(func.avg(UserOffer.cv_match_score)).filter(
                UserOffer.user_id == user_id,
                UserOffer.cv_match_score.isnot(None),
            ).scalar()
            avg_cv_score = round(float(avg_row or 0), 1)
            high_score_count = db.query(func.count(UserOffer.id)).filter(
                UserOffer.user_id == user_id,
                UserOffer.cv_match_score >= 70,
            ).scalar() or 0

        stats_data = {
            'total_offers': total_offers,
            'tracked': tracked_offers,
            'cv_sent': cv_sent,
            'follow_ups': follow_ups,
            'status_counts': status_counts,
            'response_rate': response_rate,
            'avg_cv_score': avg_cv_score,
            'high_score_count': high_score_count,
        }

        # ── Chart data ────────────────────────────────────────────────
        # Offers per source (scoped by domain)
        source_rows = (
            offer_query.with_entities(Offer.source, func.count(Offer.id))
            .group_by(Offer.source)
            .order_by(func.count(Offer.id).desc())
            .all()
        )
        source_counts = {s: c for s, c in source_rows if s}

        # Top 10 companies by offer count (scoped by domain)
        company_rows = (
            offer_query.with_entities(Offer.company, func.count(Offer.id))
            .group_by(Offer.company)
            .order_by(func.count(Offer.id).desc())
            .limit(10)
            .all()
        )
        top_companies = {c: n for c, n in company_rows if c}

        # Score distribution in 10 equal buckets (0–10, 10–20, …, 90–100)
        score_rows = offer_query.with_entities(Offer.relevance_score).all()
        score_buckets = [0] * 10
        for (score,) in score_rows:
            s = float(score or 0)
            bucket = min(int(s // 10), 9)
            score_buckets[bucket] += 1

        # CV match score distribution (only when a CV has been uploaded)
        has_cv = _has_cv_file(user_id)
        cv_score_buckets = [0] * 10
        if has_cv:
            if user_id is not None:
                # DB users: scores live in UserOffer.cv_match_score
                cv_rows = (
                    db.query(UserOffer.cv_match_score)
                    .filter(
                        UserOffer.user_id == user_id,
                        UserOffer.cv_match_score.isnot(None),
                    )
                    .all()
                )
            else:
                # Legacy admin: scores live in Offer.cv_match_score
                cv_rows = offer_query.with_entities(Offer.cv_match_score).filter(
                    Offer.cv_match_score.isnot(None)
                ).all()
            for (score,) in cv_rows:
                s = float(score or 0)
                bucket = min(int(s // 10), 9)
                cv_score_buckets[bucket] += 1

        # Weekly application timeline (date_sent grouped by ISO week)
        weekly_data = {}
        if user_id is not None:
            date_col = UserOffer.date_sent
            weekly_rows = db.query(UserOffer.date_sent).filter(
                UserOffer.user_id == user_id,
                UserOffer.date_sent.isnot(None),
            ).all()
        else:
            weekly_rows = db.query(Tracking.date_sent).filter(
                Tracking.date_sent.isnot(None),
            ).all()
        for (dt,) in weekly_rows:
            if dt:
                week_key = dt.strftime('%Y-W%W')
                weekly_data[week_key] = weekly_data.get(week_key, 0) + 1
        # Sort by week and keep last 12 weeks max
        weekly_sorted = sorted(weekly_data.items())[-12:]
        weekly_labels = [w[0] for w in weekly_sorted]
        weekly_values = [w[1] for w in weekly_sorted]

        # Contract type distribution
        contract_counts = {}
        for o in offer_query.with_entities(Offer.contract_type).all():
            ct = (o[0] or '').strip().lower()
            if 'cdi' in ct:
                key = 'CDI'
            elif 'cdd' in ct:
                key = 'CDD'
            elif 'alternance' in ct or 'apprenti' in ct:
                key = 'Alternance'
            elif 'stage' in ct:
                key = 'Stage'
            elif ct:
                key = 'Autre'
            else:
                key = 'Non précisé'
            contract_counts[key] = contract_counts.get(key, 0) + 1

        chart_data = {
            'sources':         source_counts,
            'companies':       top_companies,
            'scores':          score_buckets,
            'statuses':        status_counts,
            'cv_scores':       cv_score_buckets,
            'weekly_labels':   weekly_labels,
            'weekly_values':   weekly_values,
            'contracts':       contract_counts,
        }

        return render_template(
            'stats.html',
            stats=stats_data,
            chart_data=chart_data,
            has_cv=has_cv,
            role=get_current_role(),
            username=session.get("username"),
        )
    finally:
        db.close()


def _persist_scores(db, user_id, offers, scores: dict) -> None:
    """
    Write cv_match_score values to the database.
    DB users  → upsert into UserOffer.cv_match_score (per-user).
    Legacy admin → write into Offer.cv_match_score.
    """
    if user_id is not None:
        offer_ids = list(scores.keys())
        existing_uos = {
            uo.offer_id: uo
            for uo in db.query(UserOffer).filter(
                UserOffer.user_id == user_id,
                UserOffer.offer_id.in_(offer_ids),
            ).all()
        }
        for offer_id, score in scores.items():
            if offer_id in existing_uos:
                existing_uos[offer_id].cv_match_score = score
            else:
                db.add(UserOffer(
                    user_id=user_id,
                    offer_id=offer_id,
                    cv_match_score=score,
                    status='New',
                ))
    else:
        offer_map = {o.id: o for o in offers}
        for offer_id, score in scores.items():
            if offer_id in offer_map:
                offer_map[offer_id].cv_match_score = score


def _build_offer_query(db, domain_id, user_id, force):
    """
    Build the SQLAlchemy query for offers that need scoring.
    Returns (query, skipped_count).
    """
    query = db.query(Offer)
    if domain_id:
        query = query.filter(Offer.domain_id == domain_id)

    if user_id is not None:
        if not force:
            already_sq = (
                db.query(UserOffer.offer_id)
                .filter(
                    UserOffer.user_id == user_id,
                    UserOffer.cv_match_score.isnot(None),
                )
                .subquery()
            )
            skipped = query.filter(Offer.id.in_(already_sq)).count()
            query = query.filter(~Offer.id.in_(already_sq))
        else:
            skipped = 0
    else:
        if not force:
            skipped = query.filter(Offer.cv_match_score.isnot(None)).count()
            query = query.filter(Offer.cv_match_score.is_(None))
        else:
            skipped = 0

    return query, skipped


def _cv_matching_worker(user_id, domain_id, method: str, force: bool) -> None:
    """
    Thread worker: runs CV matching and writes scores to DB.
    Persists progress to _TASKS_FILE after each batch so all workers can poll it.
    """
    db = SessionLocal()
    try:
        cv_text = _find_cv_text(user_id)
        if not cv_text:
            _save_task(user_id, {
                'status': 'error', 'scored': 0, 'total': 0, 'skipped': 0,
                'progress': 0, 'progress_total': 0,
                'error': 'Aucun CV trouvé. Importez votre CV dans la section Documents.',
            })
            return
        query, skipped = _build_offer_query(db, domain_id, user_id, force)
        offers = query.all()
        total = len(offers)
        _save_task(user_id, {
            'status': 'running', 'scored': 0, 'total': total,
            'skipped': skipped, 'progress': 0, 'progress_total': 0, 'error': None,
        })

        if not offers:
            _save_task(user_id, {
                'status': 'done', 'scored': 0, 'total': 0,
                'skipped': skipped, 'progress': 0, 'progress_total': 0, 'error': None,
            })
            return

        if method == 'claude':
            from app.services.cv_matcher_claude import ClaudeCVMatcher
            matcher = ClaudeCVMatcher(cv_text)

            def _progress_cb(batches_done, total_batches, offers_done):
                _save_task(user_id, {
                    'status': 'running', 'scored': offers_done, 'total': total,
                    'skipped': skipped, 'progress': batches_done,
                    'progress_total': total_batches, 'error': None,
                })

            scores = matcher.score_offers(offers, progress_callback=_progress_cb)
        else:
            from app.services.cv_matcher import CVMatcher
            matcher = CVMatcher(cv_text)
            scores = matcher.score_offers(offers)

        _persist_scores(db, user_id, offers, scores)

        # Track usage counters
        if user_id is not None:
            tokens = getattr(matcher, 'total_tokens_used', 0)
            user_obj = db.query(User).filter(User.id == user_id).first()
            if user_obj:
                user_obj.matching_count = (user_obj.matching_count or 0) + 1
                if method == 'claude' and tokens > 0:
                    user_obj.claude_tokens_used = (user_obj.claude_tokens_used or 0) + tokens

        db.commit()
        _save_task(user_id, {
            'status': 'done', 'scored': len(scores), 'total': total,
            'skipped': skipped, 'progress': 0, 'progress_total': 0, 'error': None,
        })

    except Exception as e:
        logger.error(f"[cv_matching_worker] Error: {e}", exc_info=True)
        db.rollback()
        _save_task(user_id, {
            'status': 'error', 'scored': 0, 'total': 0, 'skipped': 0,
            'progress': 0, 'progress_total': 0, 'error': 'Erreur interne du serveur',
        })
    finally:
        db.close()


def _run_cv_matching(method='tfidf', force=False, domain_id=None, user_id=None):
    """
    Synchronous CV matching — used by cv_upload (tfidf, force=True).
    Returns (scored, skipped).
    """
    cv_text = _find_cv_text(user_id)
    if not cv_text:
        return 0, 0
    db = SessionLocal()
    try:
        query, skipped = _build_offer_query(db, domain_id, user_id, force)
        offers = query.all()
        if not offers:
            return 0, skipped

        if method == 'claude':
            from app.services.cv_matcher_claude import ClaudeCVMatcher
            matcher = ClaudeCVMatcher(cv_text)
        else:
            from app.services.cv_matcher import CVMatcher
            matcher = CVMatcher(cv_text)

        scores = matcher.score_offers(offers)
        _persist_scores(db, user_id, offers, scores)
        db.commit()
        return len(scores), skipped
    except Exception as e:
        db.rollback()
        raise e
    finally:
        db.close()


@bp.route('/api/cv/upload', methods=['POST'])
@admin_required
@limiter.limit("10 per minute")
def cv_upload():
    """
    Accept a PDF or plain-text CV file, extract text, save to disk,
    then run CV matching against all stored offers.
    Query param: ?method=tfidf (default) or ?method=claude
    """
    if 'cv' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['cv']
    if not file.filename:
        return jsonify({'error': 'Empty filename'}), 400

    # Extension + MIME validation for CV uploads
    _cv_allowed = {'.pdf', '.txt'}
    _cv_mimes = {
        '.pdf': {'application/pdf', 'application/octet-stream'},
        '.txt': {'text/plain', 'application/octet-stream'},
    }
    sanitized_cv_name = secure_filename(file.filename)
    if not sanitized_cv_name:
        return jsonify({'error': 'Nom de fichier invalide'}), 400
    cv_ext = Path(sanitized_cv_name).suffix.lower()
    if cv_ext not in _cv_allowed:
        return jsonify({'error': 'Seuls les fichiers .pdf et .txt sont acceptés pour le CV'}), 400
    cv_ct = (file.content_type or '').split(';')[0].strip().lower()
    if cv_ct and cv_ct not in _cv_mimes[cv_ext] and cv_ct not in _BLOCKED_MIMES:
        return jsonify({'error': f'Type MIME {cv_ct!r} incompatible avec {cv_ext}'}), 400
    if cv_ct in _BLOCKED_MIMES:
        return jsonify({'error': f'Type MIME refusé : {cv_ct}'}), 400
    # Size limit: 5 MB
    file.seek(0, 2)
    if file.tell() > MAX_UPLOAD_SIZE:
        return jsonify({'error': 'Fichier trop volumineux (max 5 Mo)'}), 400
    file.seek(0)

    filename = sanitized_cv_name.lower()
    raw = file.read()

    # Magic bytes validation
    if not _check_magic_bytes(raw, cv_ext):
        return jsonify({'error': f'Le contenu du fichier ne correspond pas à l\'extension {cv_ext}'}), 400

    # Extract text
    if cv_ext == '.pdf':
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(raw))
            cv_text = "\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        except ImportError:
            return jsonify({'error': 'PyPDF2 not installed. pip install PyPDF2'}), 500
        except Exception:
            return jsonify({'error': 'Impossible de lire le fichier PDF'}), 400
    else:
        # Assume plain text (UTF-8)
        try:
            cv_text = raw.decode('utf-8')
        except UnicodeDecodeError:
            cv_text = raw.decode('latin-1', errors='replace')

    if not cv_text.strip():
        return jsonify({'error': 'Could not extract text from CV'}), 400

    # Save to user-specific document directory
    user_id = session.get("user_id")
    if user_id is not None:
        user_cv_dir = DATA_DIR / "documents" / str(user_id)
        user_cv_dir.mkdir(parents=True, exist_ok=True)
        (user_cv_dir / filename).write_bytes(raw)
    else:
        CV_DIR.mkdir(parents=True, exist_ok=True)
        CV_TEXT_PATH.write_text(cv_text, encoding='utf-8')

    method = request.args.get('method', 'tfidf')
    if method not in ('tfidf', 'claude'):
        method = 'tfidf'

    # Run matching — force=True because a new CV invalidates all existing scores
    try:
        scored, skipped = _run_cv_matching(
            method=method,
            force=True,
            domain_id=session.get("domain_id"),
            user_id=session.get("user_id"),
        )
        return jsonify({'ok': True, 'scored': scored, 'method': method})
    except Exception as e:
        return jsonify({'error': 'Erreur interne du serveur'}), 500


@bp.route('/api/cv/rematch', methods=['POST'])
@admin_required
@limiter.limit("5 per minute")
def cv_rematch():
    """
    Launch CV matching asynchronously in a background thread and return immediately.

    Query params:
      method=tfidf|claude  — scoring engine (default: tfidf)
      force=true           — re-score ALL offers, even those already scored

    Returns immediately with {ok, status: 'started'|'already_running', task_id}.
    Poll GET /api/cv/matching-status for progress.
    """
    user_id = session.get('user_id')
    if not _has_cv_file(user_id):
        return jsonify({'error': 'Aucun CV trouvé. Importez votre CV dans la section Documents.'}), 404

    method = request.args.get('method', 'tfidf')
    if method not in ('tfidf', 'claude'):
        method = 'tfidf'
    force = request.args.get('force', 'false').lower() == 'true'

    domain_id = session.get('domain_id')

    # Quota check — only for claude method on role=user accounts
    if method == 'claude':
        allowed, used, limit = _check_and_increment_quota(
            user_id, 'weekly_matches_used', WEEKLY_MATCH_LIMIT
        )
        if not allowed:
            return jsonify({
                'error': f'Quota atteint : {limit} matching(s) IA par semaine. Réessayez lundi.',
                'quota_exceeded': True,
                'used': used,
                'limit': limit,
            }), 429

    task_id = str(uuid.uuid4())
    initial_state = {
        'status': 'running',
        'task_id': task_id,
        'scored': 0,
        'total': 0,
        'skipped': 0,
        'progress': 0,
        'progress_total': 0,
        'error': None,
    }

    if not _try_start_task(user_id, initial_state):
        existing = _load_task(user_id) or {}
        return jsonify({
            'ok': False,
            'status': 'already_running',
            'scored': existing.get('scored', 0),
            'total': existing.get('total', 0),
        })

    thread = threading.Thread(
        target=_cv_matching_worker,
        args=(user_id, domain_id, method, force),
        daemon=True,
    )
    thread.start()

    return jsonify({'ok': True, 'status': 'started', 'task_id': task_id})


@bp.route('/api/cv/matching-status')
@login_required
def cv_matching_status():
    """Return the current CV matching task state for the logged-in user."""
    user_id = session.get('user_id')
    task = _load_task(user_id)
    if not task:
        return jsonify({'ok': True, 'status': 'none'})
    return jsonify({
        'ok': True,
        'status':          task['status'],
        'scored':          task.get('scored', 0),
        'total':           task.get('total', 0),
        'skipped':         task.get('skipped', 0),
        'progress':        task.get('progress', 0),
        'progress_total':  task.get('progress_total', 0),
        'error':           task.get('error'),
    })


# ── Account management (password change + TOTP 2FA) ───────────────────────────

def _qr_code_b64(uri: str) -> str:
    """Generate a base64-encoded PNG QR code for the given provisioning URI."""
    import qrcode as _qrcode
    import base64
    qr = _qrcode.QRCode(box_size=5, border=2,
                        error_correction=_qrcode.constants.ERROR_CORRECT_L)
    qr.add_data(uri)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    return base64.b64encode(buf.getvalue()).decode()


@bp.route('/account', methods=['GET', 'POST'])
@login_required
def account():
    """Account page: password change and 2FA management."""
    user_id = session.get('user_id')
    errors = []
    success = None

    # Handle password change (POST with action=change_password)
    if request.method == 'POST' and request.form.get('action') == 'change_password':
        if user_id is None:
            errors.append("Changement de mot de passe non disponible pour les comptes de configuration.")
        else:
            from app import bcrypt
            current_pw = request.form.get('current_password', '')
            new_pw = request.form.get('new_password', '')
            confirm_pw = request.form.get('confirm_password', '')

            db = SessionLocal()
            try:
                user = db.query(User).filter(User.id == user_id).first()
                if not user or not bcrypt.check_password_hash(user.password_hash, current_pw):
                    errors.append("Mot de passe actuel incorrect.")
                else:
                    errors.extend(_validate_password(new_pw))
                    if new_pw != confirm_pw:
                        errors.append("Les nouveaux mots de passe ne correspondent pas.")
                if not errors:
                    user.password_hash = bcrypt.generate_password_hash(new_pw).decode('utf-8')
                    user.updated_at = datetime.utcnow()
                    db.commit()
                    # Regenerate session to invalidate old session cookies
                    _uname = session.get("username")
                    _role = session.get("role")
                    _uid = session.get("user_id")
                    _did = session.get("domain_id")
                    session.clear()
                    session["username"] = _uname
                    session["role"] = _role
                    session["user_id"] = _uid
                    session["domain_id"] = _did
                    success = "Mot de passe modifié avec succès."
            finally:
                db.close()

    # Resolve current 2FA state
    totp_enabled = False
    has_db_user = user_id is not None
    qr_b64 = None
    qr_uri = None

    if has_db_user:
        db = SessionLocal()
        try:
            user = db.query(User).filter(User.id == user_id).first()
            totp_enabled = bool(user and user.totp_enabled)
        finally:
            db.close()

    # Show QR code if a setup is in progress
    if has_db_user and session.get('_totp_setup_secret') and not totp_enabled:
        import pyotp
        secret = session['_totp_setup_secret']
        uri = pyotp.TOTP(secret).provisioning_uri(
            name=session.get('username', ''), issuer_name='MyJobHunter'
        )
        try:
            qr_b64 = _qr_code_b64(uri)
            qr_uri = uri
        except Exception:
            qr_b64 = None

    return render_template(
        'account.html',
        errors=errors,
        success=success,
        has_db_user=has_db_user,
        totp_enabled=totp_enabled,
        qr_b64=qr_b64,
        totp_setup_secret=session.get('_totp_setup_secret') if qr_b64 else None,
        role=get_current_role(),
        username=session.get('username'),
    )


@bp.route('/account/setup-2fa', methods=['POST'])
@login_required
def account_setup_2fa():
    """Generate a new TOTP secret and store it in session for confirmation."""
    user_id = session.get('user_id')
    if user_id is None:
        return redirect(url_for('main.account'))
    import pyotp
    secret = pyotp.random_base32()
    session['_totp_setup_secret'] = secret
    return redirect(url_for('main.account'))


@bp.route('/account/confirm-2fa', methods=['POST'])
@login_required
def account_confirm_2fa():
    """Verify TOTP code and activate 2FA for the user."""
    user_id = session.get('user_id')
    secret = session.get('_totp_setup_secret')
    if not user_id or not secret:
        return redirect(url_for('main.account'))

    import pyotp
    code = request.form.get('totp_code', '').strip().replace(' ', '')
    totp = pyotp.TOTP(secret)
    if not totp.verify(code, valid_window=0):
        # Keep setup in session, show error via flash-like query param
        return redirect(url_for('main.account') + '?err=invalid_code')

    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if user:
            user.totp_secret = _encrypt_totp_secret(secret)
            user.totp_enabled = True
            user.updated_at = datetime.utcnow()
            db.commit()
            _sec_log("2FA_ENABLED", user.username)
    finally:
        db.close()

    session.pop('_totp_setup_secret', None)
    return redirect(url_for('main.account') + '?ok=2fa_enabled')


@bp.route('/account/disable-2fa', methods=['POST'])
@login_required
def account_disable_2fa():
    """Verify current TOTP code and disable 2FA for the user."""
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('main.account'))

    import pyotp
    code = request.form.get('totp_code', '').strip().replace(' ', '')

    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if not user or not user.totp_enabled or not user.totp_secret:
            return redirect(url_for('main.account'))
        totp = pyotp.TOTP(_decrypt_totp_secret(user.totp_secret))
        if not totp.verify(code, valid_window=0):
            return redirect(url_for('main.account') + '?err=invalid_code')
        _sec_log("2FA_DISABLED", user.username)
        user.totp_secret = None
        user.totp_enabled = False
        user.updated_at = datetime.utcnow()
        db.commit()
    finally:
        db.close()

    return redirect(url_for('main.account') + '?ok=2fa_disabled')


# ── Account profile ───────────────────────────────────────────────────────────

@bp.route('/account/profile', methods=['GET', 'POST'])
@login_required
def account_profile():
    """Profile page: info display, domain change, email change."""
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('main.account'))

    errors = []
    success = None

    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            return redirect(url_for('main.logout'))

        if request.method == 'POST':
            action = request.form.get('action')

            if action == 'change_domain':
                new_domain_id = request.form.get('domain_id', '').strip()
                if new_domain_id == '':
                    new_domain_id = None
                else:
                    try:
                        new_domain_id = int(new_domain_id)
                    except ValueError:
                        errors.append("Domaine invalide.")
                        new_domain_id = user.domain_id

                if not errors and new_domain_id != user.domain_id:
                    # Reset cv_match_score on all user_offers
                    db.query(UserOffer).filter(UserOffer.user_id == user_id).update(
                        {UserOffer.cv_match_score: None},
                        synchronize_session='fetch'
                    )
                    user.domain_id = new_domain_id
                    user.updated_at = datetime.utcnow()
                    db.commit()
                    session['domain_id'] = new_domain_id
                    success = "Domaine mis à jour. Les scores Match IA ont été réinitialisés."

            elif action == 'change_email':
                new_email = request.form.get('email', '').strip()
                import re as _re
                _email_re = _re.compile(r'^[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}$')
                if new_email and not _email_re.match(new_email):
                    errors.append("Adresse e-mail invalide.")
                else:
                    user.email = new_email or None
                    user.updated_at = datetime.utcnow()
                    db.commit()
                    success = "Adresse e-mail mise à jour."

        # Reload user after possible commit
        db.refresh(user)
        domains = db.query(Domain).order_by(Domain.name).all()

        return render_template(
            'profile.html',
            user=user,
            domains=domains,
            errors=errors,
            success=success,
            role=get_current_role(),
            username=session.get('username'),
        )
    finally:
        db.close()


@bp.route('/api/account/delete', methods=['POST'])
@login_required
def account_delete():
    """Self-service account deletion — requires typing 'SUPPRIMER' to confirm."""
    import shutil

    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'ok': False, 'error': 'Non autorisé'}), 403

    confirm = request.form.get('confirm_text', '').strip()
    if confirm != 'SUPPRIMER':
        return jsonify({'ok': False, 'error': 'Confirmation incorrecte'}), 400

    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            return jsonify({'ok': False, 'error': 'Utilisateur introuvable'}), 404

        # Delete documents directory
        docs_dir = Path(DATA_DIR) / 'documents' / str(user_id)
        if docs_dir.exists():
            shutil.rmtree(docs_dir, ignore_errors=True)

        # Delete user (cascades: user_offers, password_resets)
        _sec_log("ACCOUNT_DELETE", user.username, "self-delete")
        db.delete(user)
        db.commit()
    finally:
        db.close()

    session.clear()
    return jsonify({'ok': True})


# ── Document management ───────────────────────────────────────────────────────

@bp.route('/documents')
@login_required
def documents():
    """Document library: list uploaded files (CV, cover letters, etc.)."""
    docs_dir = _user_docs_dir()
    docs_dir.mkdir(parents=True, exist_ok=True)
    files = sorted(f.name for f in docs_dir.iterdir() if f.is_file())
    return render_template('documents.html', files=files,
                           role=get_current_role(),
                           username=session.get("username"))


@bp.route('/api/documents/upload', methods=['POST'])
@admin_required
@limiter.limit("10 per minute")
def document_upload():
    """Upload a document file (PDF, TXT, DOCX) with strict validation."""
    if 'file' not in request.files:
        return jsonify({'error': 'Aucun fichier fourni'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': 'Nom de fichier vide'}), 400

    # 1. Sanitize filename
    filename = secure_filename(file.filename)
    if not filename:
        return jsonify({'error': 'Nom de fichier invalide après sanitisation'}), 400

    # 2. Extension check — must have exactly one recognised extension
    p = Path(filename)
    ext = p.suffix.lower()
    if not ext:
        return jsonify({'error': 'Le fichier doit avoir une extension (.pdf, .docx, .txt)'}), 400
    if ext not in _ALLOWED_MIMES:
        return jsonify({'error': f'Extension refusée : {ext}. Acceptés : .pdf, .docx, .txt'}), 400

    # 3. Double-extension check (e.g. malware.exe.pdf)
    stem_ext = Path(p.stem).suffix.lower()
    if stem_ext:
        return jsonify({'error': 'Double extension refusée (ex: fichier.exe.pdf)'}), 400

    # 4. MIME type validation
    content_type = (file.content_type or '').split(';')[0].strip().lower()
    if content_type in _BLOCKED_MIMES:
        return jsonify({'error': f'Type MIME refusé : {content_type}'}), 400
    if content_type and content_type not in _ALLOWED_MIMES[ext]:
        return jsonify({'error': f'Type MIME {content_type!r} incompatible avec {ext}'}), 400

    # 5. Size check (read into memory to measure; limit stream)
    file.seek(0, 2)
    size = file.tell()
    file.seek(0)
    if size > MAX_UPLOAD_SIZE:
        mb = size / (1024 * 1024)
        return jsonify({'error': f'Fichier trop volumineux ({mb:.1f} Mo). Max : 5 Mo'}), 400

    # 6. Magic bytes validation
    raw = file.read()
    if not _check_magic_bytes(raw, ext):
        return jsonify({'error': f'Le contenu du fichier ne correspond pas à l\'extension {ext}'}), 400

    docs_dir = _user_docs_dir()
    docs_dir.mkdir(parents=True, exist_ok=True)
    (docs_dir / filename).write_bytes(raw)
    return jsonify({'ok': True, 'filename': filename})


@bp.route('/api/documents/<filename>', methods=['GET'])
@login_required
def document_download(filename):
    """Download an uploaded document."""
    base_dir = _user_docs_dir().resolve()
    filepath = (base_dir / secure_filename(filename)).resolve()
    if not str(filepath).startswith(str(base_dir)):
        return "Accès refusé", 403
    if not filepath.exists():
        return "Fichier introuvable", 404
    return send_file(str(filepath), as_attachment=True,
                     download_name=filepath.name)


@bp.route('/api/documents/<filename>', methods=['DELETE'])
@login_required
def document_delete(filename):
    """Delete an uploaded document."""
    base_dir = _user_docs_dir().resolve()
    filepath = (base_dir / secure_filename(filename)).resolve()
    if not str(filepath).startswith(str(base_dir)):
        return jsonify({'error': 'Accès refusé'}), 403
    if not filepath.exists():
        return jsonify({'error': 'Fichier introuvable'}), 404
    filepath.unlink()
    return jsonify({'ok': True})


# ── Cover letter generation ───────────────────────────────────────────────────

@bp.route('/api/cover-letter/<int:offer_id>', methods=['POST'])
@admin_required
@limiter.limit("5 per minute")
def generate_cover_letter(offer_id):
    """
    Generate a tailored cover letter for an offer using Claude.
    Body JSON: { template_filename: str|"", format: "txt"|"docx" }
    Returns JSON { ok, text, filename } for txt, or binary for docx.
    """
    from config import APIKeys

    db = SessionLocal()
    try:
        offer = db.query(Offer).filter(Offer.id == offer_id).first()
        if not offer:
            return jsonify({'error': 'Offre introuvable'}), 404

        # Domain authorization: users may only act on offers in their domain
        _domain_id = session.get('domain_id')
        if _domain_id and offer.domain_id and offer.domain_id != _domain_id:
            return jsonify({'error': 'Accès refusé'}), 403

        # Quota check for cover letter generation
        _uid = session.get('user_id')
        _allowed, _used, _limit = _check_and_increment_quota(
            _uid, 'weekly_letters_used', WEEKLY_LETTER_LIMIT
        )
        if not _allowed:
            return jsonify({
                'error': f'Quota atteint : {_limit} lettre(s) de motivation par semaine. Réessayez lundi.',
                'quota_exceeded': True,
                'used': _used,
                'limit': _limit,
            }), 429

        data = request.get_json() or {}
        template_filename = data.get('template_filename', '').strip()
        output_format = data.get('format', 'txt')
        if output_format not in ('txt', 'docx'):
            output_format = 'txt'

        # Read cover letter template (if provided)
        template_text = ""
        if template_filename:
            tpl_path = _user_docs_dir() / secure_filename(template_filename)
            if tpl_path.exists():
                ext = tpl_path.suffix.lower()
                if ext == '.pdf':
                    try:
                        import PyPDF2
                        reader = PyPDF2.PdfReader(str(tpl_path))
                        template_text = "\n".join(
                            p.extract_text() or "" for p in reader.pages
                        )
                    except Exception:
                        pass
                else:
                    try:
                        template_text = tpl_path.read_text(encoding='utf-8')
                    except UnicodeDecodeError:
                        template_text = tpl_path.read_text(
                            encoding='latin-1', errors='replace'
                        )

        # Build prompt
        if template_text.strip():
            intro = (
                "Adapte cette lettre de motivation existante à la nouvelle offre :\n\n"
                "---\n" + template_text.strip() + "\n---\n\n"
            )
        else:
            intro = "Rédige une lettre de motivation professionnelle.\n\n"

        prompt = (
            intro
            + "Offre ciblée :\n"
            + f"- Poste : {offer.title}\n"
            + f"- Entreprise : {offer.company}\n"
            + f"- Description : {(offer.description or '')[:2000]}\n\n"
            + "Instructions :\n"
            + "- Personnalise l'introduction en mentionnant l'entreprise et le poste\n"
            + "- Mets en avant les compétences les plus pertinentes pour ce rôle\n"
            + "- Ton professionnel et enthousiaste, 3 à 4 paragraphes\n"
            + "- Réponds UNIQUEMENT avec la lettre, sans titre ni commentaires"
        )

        if not APIKeys.ANTHROPIC_API_KEY:
            return jsonify({'error': 'ANTHROPIC_API_KEY non configurée dans .env'}), 503

        from anthropic import Anthropic
        client = Anthropic(api_key=APIKeys.ANTHROPIC_API_KEY)
        message = client.messages.create(
            model=APIKeys.ANTHROPIC_MODEL,
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}],
        )
        letter_text = message.content[0].text
        if hasattr(message, 'usage') and message.usage:
            _add_claude_tokens(
                session.get('user_id'),
                (message.usage.input_tokens or 0) + (message.usage.output_tokens or 0),
            )

        # Sanitize company name for filename
        company_safe = "".join(
            c if c.isalnum() or c in " _-" else "_"
            for c in offer.company
        ).strip().replace(" ", "_")

        # Return binary .docx if python-docx is available
        if output_format == 'docx':
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument()
                for para in letter_text.split('\n\n'):
                    if para.strip():
                        doc.add_paragraph(para.strip())
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                dl_name = f"lettre_{company_safe}.docx"
                mime = (
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document"
                )
                return send_file(buf, as_attachment=True,
                                 download_name=dl_name, mimetype=mime)
            except ImportError:
                pass  # fall through to txt

        filename = f"lettre_{company_safe}.txt"
        return jsonify({'ok': True, 'text': letter_text, 'filename': filename})

    except Exception as e:
        return jsonify({'error': 'Erreur interne du serveur'}), 500
    finally:
        db.close()


# ── Admin panel ───────────────────────────────────────────────────────────────

@bp.route('/admin')
@superadmin_required
@limiter.limit("30 per minute")
def admin_page():
    """Admin panel: list all registered users with management actions."""
    db = SessionLocal()
    try:
        users = db.query(User).order_by(User.created_at.desc()).all()
        domains = {d.id: d.name for d in db.query(Domain).all()}
        # Count documents per user
        doc_counts: dict[int, int] = {}
        for u in users:
            d = _admin_user_docs_dir(u.id)
            doc_counts[u.id] = sum(1 for f in d.iterdir() if f.is_file()) if d.exists() else 0

        # Count tracked offers per user
        offer_counts: dict[int, int] = {
            uid: cnt
            for uid, cnt in db.query(UserOffer.user_id, func.count(UserOffer.id))
                               .group_by(UserOffer.user_id)
                               .all()
        }

        return render_template(
            'admin.html',
            users=users,
            domains=domains,
            doc_counts=doc_counts,
            offer_counts=offer_counts,
            now=datetime.utcnow(),
            role=get_current_role(),
            username=session.get("username"),
        )
    finally:
        db.close()


@bp.route('/api/admin/users/<int:user_id>/toggle', methods=['POST'])
@superadmin_required
@limiter.limit("30 per minute")
def admin_toggle_user(user_id):
    """Enable or disable a user account."""
    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            return jsonify({'error': 'User not found'}), 404
        if user.username == session.get("username"):
            return jsonify({'error': 'Impossible de désactiver votre propre compte'}), 400
        user.is_active = not user.is_active
        # When admin manually activates, also mark email as confirmed
        if user.is_active:
            user.email_confirmed = True
        db.commit()
        return jsonify({'ok': True, 'is_active': user.is_active})
    except Exception as e:
        db.rollback()
        return jsonify({'error': 'Erreur interne du serveur'}), 500
    finally:
        db.close()


@bp.route('/api/admin/users/<int:user_id>/delete', methods=['POST'])
@superadmin_required
@limiter.limit("30 per minute")
def admin_delete_user(user_id):
    """Delete a user account and all associated data."""
    if user_id == session.get('user_id'):
        return jsonify({'error': 'Impossible de supprimer votre propre compte'}), 400
    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            return jsonify({'error': 'Utilisateur introuvable'}), 404
        # Delete documents directory
        import shutil
        docs_dir = _admin_user_docs_dir(user_id)
        if docs_dir.exists():
            shutil.rmtree(str(docs_dir))
        # Cascade deletes user_offers and password_resets via FK cascade
        _sec_log("ACCOUNT_DELETE", user.username, f"by_admin={session.get('username')}")
        db.delete(user)
        db.commit()
        return jsonify({'ok': True})
    except Exception:
        db.rollback()
        return jsonify({'error': 'Erreur interne du serveur'}), 500
    finally:
        db.close()


# ── Admin: per-user document management ───────────────────────────────────────

def _admin_user_docs_dir(user_id: int) -> Path:
    """Return the document directory for a given user (admin access)."""
    return DATA_DIR / "documents" / str(user_id)


@bp.route('/admin/documents/<int:user_id>')
@superadmin_required
@limiter.limit("30 per minute")
def admin_user_documents(user_id):
    """Admin view: list and manage documents belonging to a specific user."""
    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            return "Utilisateur introuvable", 404
        docs_dir = _admin_user_docs_dir(user_id)
        files = sorted(f.name for f in docs_dir.iterdir() if f.is_file()) if docs_dir.exists() else []
        return render_template(
            'admin_user_docs.html',
            target_user=user,
            files=files,
            role=get_current_role(),
            username=session.get("username"),
        )
    finally:
        db.close()


@bp.route('/api/admin/documents/<int:user_id>/<filename>', methods=['GET'])
@superadmin_required
@limiter.limit("30 per minute")
def admin_document_download(user_id, filename):
    """Admin download of a specific user's document."""
    base_dir = _admin_user_docs_dir(user_id).resolve()
    filepath = (base_dir / secure_filename(filename)).resolve()
    if not str(filepath).startswith(str(base_dir)):
        return "Accès refusé", 403
    if not filepath.exists():
        return "Fichier introuvable", 404
    return send_file(str(filepath), as_attachment=True, download_name=filepath.name)


@bp.route('/api/admin/documents/<int:user_id>/<filename>', methods=['DELETE'])
@superadmin_required
@limiter.limit("30 per minute")
def admin_document_delete(user_id, filename):
    """Admin delete of a specific user's document."""
    base_dir = _admin_user_docs_dir(user_id).resolve()
    filepath = (base_dir / secure_filename(filename)).resolve()
    if not str(filepath).startswith(str(base_dir)):
        return jsonify({'error': 'Accès refusé'}), 403
    if not filepath.exists():
        return jsonify({'error': 'Fichier introuvable'}), 404
    filepath.unlink()
    return jsonify({'ok': True})


# ── Password reset ─────────────────────────────────────────────────────────────

@bp.route('/api/admin/users/<int:user_id>/reset-password', methods=['POST'])
@superadmin_required
@limiter.limit("30 per minute")
def admin_reset_password(user_id):
    """Generate a single-use 15-min password reset link for a user (admin only)."""
    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            return jsonify({'error': 'Utilisateur introuvable'}), 404
        # Invalidate any existing unused tokens for this user
        db.query(PasswordReset).filter(
            PasswordReset.user_id == user_id,
            PasswordReset.used == False
        ).update({'used': True})
        token = secrets.token_urlsafe(32)
        db.add(PasswordReset(user_id=user_id, token=token))
        db.commit()
        reset_url = url_for('main.reset_password', token=token, _external=True, _scheme='https')
        return jsonify({'ok': True, 'reset_url': reset_url})
    except Exception:
        db.rollback()
        return jsonify({'error': 'Erreur interne'}), 500
    finally:
        db.close()


@bp.route('/forgot-password', methods=['GET', 'POST'])
@limiter.limit("10 per minute")
def forgot_password():
    """Account recovery: email link if user has email, else security question flow."""
    if request.method == 'GET':
        return render_template('forgot_password.html', step='1',
                               security_questions=SECURITY_QUESTIONS)

    action = request.form.get('action', '')

    if action == 'get_question':
        username = request.form.get('username', '').strip()
        db = SessionLocal()
        try:
            user = db.query(User).filter(
                User.username == username, User.is_active == True
            ).first()
            if not user:
                # Return same response as email-sent to prevent username enumeration
                return render_template('forgot_password.html', step='email_sent')
            # ── Email path: user has an email → send reset link directly ──────
            if user.email:
                db.query(PasswordReset).filter(
                    PasswordReset.user_id == user.id,
                    PasswordReset.used == False,
                ).update({'used': True})
                token = secrets.token_urlsafe(32)
                db.add(PasswordReset(user_id=user.id, token=token))
                db.commit()
                reset_url = url_for('main.reset_password', token=token,
                                    _external=True, _scheme='https')
                _send_reset_email(user.email, username, reset_url)
                return render_template('forgot_password.html', step='email_sent')
            # ── Security question path ─────────────────────────────────────────
            if not user.security_question:
                return render_template(
                    'forgot_password.html', step='1',
                    error="Aucun email ni question de sécurité configurés pour ce compte. Contactez un administrateur.",
                    security_questions=SECURITY_QUESTIONS,
                )
            return render_template('forgot_password.html', step='2',
                                   username=username,
                                   question=user.security_question)
        finally:
            db.close()

    if action == 'verify_answer':
        username = request.form.get('username', '').strip()
        answer = request.form.get('answer', '').strip()
        db = SessionLocal()
        try:
            from app import bcrypt
            user = db.query(User).filter(
                User.username == username, User.is_active == True
            ).first()
            if not user or not user.security_answer_hash:
                return render_template('forgot_password.html', step='1',
                                       error="Utilisateur introuvable.",
                                       security_questions=SECURITY_QUESTIONS)
            # Check lockout
            now = datetime.utcnow()
            if user.security_lockout_until and now < user.security_lockout_until:
                remaining = int((user.security_lockout_until - now).total_seconds() / 60) + 1
                return render_template('forgot_password.html', step='2',
                                       username=username,
                                       question=user.security_question,
                                       error=f"Trop de tentatives, réessayez dans {remaining} minutes.")
            if not bcrypt.check_password_hash(user.security_answer_hash, answer.lower()):
                user.failed_security_attempts = (user.failed_security_attempts or 0) + 1
                _sec_log("SEC_QUESTION_FAIL", username, f"attempt={user.failed_security_attempts}")
                if user.failed_security_attempts >= 5:
                    user.security_lockout_until = now + timedelta(minutes=30)
                    db.commit()
                    _sec_log("SEC_QUESTION_LOCK", username, f"lockout_until={user.security_lockout_until.isoformat()}")
                    return render_template('forgot_password.html', step='2',
                                           username=username,
                                           question=user.security_question,
                                           error="Trop de tentatives, réessayez dans 30 minutes.")
                db.commit()
                return render_template('forgot_password.html', step='2',
                                       username=username,
                                       question=user.security_question,
                                       error="Réponse incorrecte. Vérifiez votre réponse.")
            # Correct answer — reset counter and generate token
            user.failed_security_attempts = 0
            user.security_lockout_until = None
            db.query(PasswordReset).filter(
                PasswordReset.user_id == user.id,
                PasswordReset.used == False
            ).update({'used': True})
            token = secrets.token_urlsafe(32)
            db.add(PasswordReset(user_id=user.id, token=token))
            db.commit()
            reset_url = url_for('main.reset_password', token=token, _external=True, _scheme='https')
            return render_template('forgot_password.html', step='3', reset_url=reset_url)
        finally:
            db.close()

    return render_template('forgot_password.html', step='1',
                           security_questions=SECURITY_QUESTIONS)


def _send_reset_email(to_email: str, username: str, reset_url: str) -> None:
    """Send a password reset email. Silently logs on failure."""
    import html as _html
    from flask_mail import Message
    from app import mail
    # Escape user-controlled values before embedding in HTML
    safe_username = _html.escape(username)
    safe_reset_url = _html.escape(reset_url)
    html_body = f"""<!DOCTYPE html>
<html lang="fr">
<head><meta charset="UTF-8"></head>
<body style="margin:0;padding:0;background:#f8fafc;font-family:'Helvetica Neue',Arial,sans-serif;">
  <table width="100%" cellpadding="0" cellspacing="0" style="background:#f8fafc;padding:40px 0;">
    <tr><td align="center">
      <table width="520" cellpadding="0" cellspacing="0"
             style="background:#ffffff;border-radius:12px;box-shadow:0 4px 24px rgba(0,0,0,.08);overflow:hidden;">
        <!-- Header -->
        <tr>
          <td style="background:#2563eb;padding:28px 40px;text-align:center;">
            <span style="font-size:2.2rem;">🎯</span>
            <h1 style="margin:8px 0 0;color:#ffffff;font-size:1.4rem;font-weight:700;letter-spacing:-.3px;">
              MyJobHunter
            </h1>
          </td>
        </tr>
        <!-- Body -->
        <tr>
          <td style="padding:36px 40px;">
            <h2 style="margin:0 0 12px;font-size:1.15rem;color:#0f172a;">
              Réinitialisation de votre mot de passe
            </h2>
            <p style="margin:0 0 16px;color:#475569;font-size:.95rem;line-height:1.6;">
              Bonjour <strong>{safe_username}</strong>,
            </p>
            <p style="margin:0 0 24px;color:#475569;font-size:.95rem;line-height:1.6;">
              Vous avez demandé à réinitialiser le mot de passe de votre compte MyJobHunter.
              Cliquez sur le bouton ci-dessous pour choisir un nouveau mot de passe.
            </p>
            <!-- CTA button -->
            <table cellpadding="0" cellspacing="0" width="100%" style="margin-bottom:24px;">
              <tr>
                <td align="center">
                  <a href="{safe_reset_url}"
                     style="display:inline-block;background:#2563eb;color:#ffffff;
                            text-decoration:none;padding:14px 36px;border-radius:8px;
                            font-size:1rem;font-weight:600;letter-spacing:-.2px;">
                    Réinitialiser mon mot de passe →
                  </a>
                </td>
              </tr>
            </table>
            <p style="margin:0 0 8px;color:#94a3b8;font-size:.82rem;line-height:1.5;">
              Ce lien est valable <strong>15 minutes</strong> et ne peut être utilisé qu'une seule fois.
            </p>
            <p style="margin:0;color:#94a3b8;font-size:.82rem;line-height:1.5;">
              Si vous n'avez pas demandé cette réinitialisation, ignorez cet email — votre mot de passe ne changera pas.
            </p>
          </td>
        </tr>
        <!-- Footer -->
        <tr>
          <td style="background:#f8fafc;padding:20px 40px;text-align:center;
                     border-top:1px solid #e2e8f0;">
            <p style="margin:0;color:#94a3b8;font-size:.78rem;">
              © 2026 MyJobHunter · Cet email est automatique, ne pas répondre.
            </p>
          </td>
        </tr>
      </table>
    </td></tr>
  </table>
</body>
</html>"""
    try:
        msg = Message(
            subject="MyJobHunter - Réinitialisation de votre mot de passe",
            recipients=[to_email],
            html=html_body,
        )
        mail.send(msg)
    except Exception as exc:
        logger.error("Failed to send reset email to %s: %s", to_email, exc)


def _send_confirmation_email(to_email: str, username: str, confirm_url: str) -> None:
    """Send an email confirmation link after registration. Silently logs on failure."""
    import html as _html
    from flask_mail import Message
    from app import mail
    safe_username = _html.escape(username)
    safe_confirm_url = _html.escape(confirm_url)
    html_body = f"""<!DOCTYPE html>
<html lang="fr">
<head><meta charset="UTF-8"></head>
<body style="margin:0;padding:0;background:#f8fafc;font-family:'Helvetica Neue',Arial,sans-serif;">
  <table width="100%" cellpadding="0" cellspacing="0" style="background:#f8fafc;padding:40px 0;">
    <tr><td align="center">
      <table width="520" cellpadding="0" cellspacing="0"
             style="background:#ffffff;border-radius:12px;box-shadow:0 4px 24px rgba(0,0,0,.08);overflow:hidden;">
        <tr>
          <td style="background:#2563eb;padding:28px 40px;text-align:center;">
            <span style="font-size:2.2rem;">&#127919;</span>
            <h1 style="margin:8px 0 0;color:#ffffff;font-size:1.4rem;font-weight:700;letter-spacing:-.3px;">
              MyJobHunter
            </h1>
          </td>
        </tr>
        <tr>
          <td style="padding:36px 40px;">
            <h2 style="margin:0 0 12px;font-size:1.15rem;color:#0f172a;">
              Confirmez votre adresse email
            </h2>
            <p style="margin:0 0 16px;color:#475569;font-size:.95rem;line-height:1.6;">
              Bonjour <strong>{safe_username}</strong>,
            </p>
            <p style="margin:0 0 24px;color:#475569;font-size:.95rem;line-height:1.6;">
              Bienvenue sur MyJobHunter ! Pour activer votre compte, veuillez confirmer
              votre adresse email en cliquant sur le bouton ci-dessous.
            </p>
            <table cellpadding="0" cellspacing="0" width="100%" style="margin-bottom:24px;">
              <tr>
                <td align="center">
                  <a href="{safe_confirm_url}"
                     style="display:inline-block;background:#2563eb;color:#ffffff;
                            text-decoration:none;padding:14px 36px;border-radius:8px;
                            font-size:1rem;font-weight:600;letter-spacing:-.2px;">
                    Confirmer mon email &#8594;
                  </a>
                </td>
              </tr>
            </table>
            <p style="margin:0 0 8px;color:#94a3b8;font-size:.82rem;line-height:1.5;">
              Ce lien est valable <strong>24 heures</strong> et ne peut être utilisé qu'une seule fois.
            </p>
            <p style="margin:0;color:#94a3b8;font-size:.82rem;line-height:1.5;">
              Si vous n'avez pas créé de compte sur MyJobHunter, ignorez cet email.
            </p>
          </td>
        </tr>
        <tr>
          <td style="background:#f8fafc;padding:20px 40px;text-align:center;
                     border-top:1px solid #e2e8f0;">
            <p style="margin:0;color:#94a3b8;font-size:.78rem;">
              &copy; 2026 MyJobHunter &middot; Cet email est automatique, ne pas répondre.
            </p>
          </td>
        </tr>
      </table>
    </td></tr>
  </table>
</body>
</html>"""
    try:
        msg = Message(
            subject="MyJobHunter - Confirmez votre adresse email",
            recipients=[to_email],
            html=html_body,
        )
        mail.send(msg)
    except Exception as exc:
        logger.error("Failed to send confirmation email to %s: %s", to_email, exc)


@bp.route('/reset/<token>', methods=['GET', 'POST'])
@limiter.limit("10 per minute")
def reset_password(token):
    """Verify a password reset token and allow the user to set a new password."""
    db = SessionLocal()
    try:
        reset = db.query(PasswordReset).filter(
            PasswordReset.token == token,
            PasswordReset.used == False,
        ).first()
        if not reset:
            return render_template('reset_password.html',
                                   error="Lien invalide ou déjà utilisé.")
        if datetime.utcnow() - reset.created_at > timedelta(minutes=15):
            reset.used = True
            db.commit()
            return render_template('reset_password.html',
                                   error="Ce lien a expiré (validité 15 minutes). Recommencez la procédure.")

        errors = []
        if request.method == 'POST':
            new_pw = request.form.get('new_password', '')
            confirm = request.form.get('confirm_password', '')
            errors.extend(_validate_password(new_pw))
            if new_pw != confirm:
                errors.append("Les mots de passe ne correspondent pas.")
            if not errors:
                from app import bcrypt
                user = db.query(User).filter(User.id == reset.user_id).first()
                if user:
                    user.password_hash = bcrypt.generate_password_hash(new_pw).decode('utf-8')
                    user.updated_at = datetime.utcnow()
                    reset.used = True
                    db.commit()
                    _sec_log("PASSWORD_RESET", user.username)
                    # Clear any active session for this user (current browser)
                    session.clear()
                    return redirect(url_for('main.login') + '?ok=password_reset')
        return render_template('reset_password.html', token=token, errors=errors)
    finally:
        db.close()


# ── Admin: site statistics ─────────────────────────────────────────────────────

_NGINX_LOG = Path('/var/log/nginx/access.log')


def _parse_nginx_today():
    """Parse today's entries from the nginx access log.
    Returns (unique_ips, total_requests, top_pages, landing_hits).
    top_pages is a list of (path, count) tuples, max 10.
    Returns None values if the log file is unavailable.
    """
    import re
    from collections import Counter
    today_prefix = datetime.now().strftime('%d/%b/%Y')
    # Match: IP ... [DD/Mon/YYYY: ... "METHOD /path HTTP ..."
    pat = re.compile(
        r'^(\S+)\s+\S+\s+\S+\s+\[' + re.escape(today_prefix) +
        r':[^\]]+\]\s+"(?:GET|POST|PUT|DELETE|HEAD|OPTIONS|PATCH)\s+(\S+)\s+HTTP'
    )
    unique_ips: set = set()
    total = 0
    counter: Counter = Counter()
    try:
        with open(_NGINX_LOG, 'r', errors='replace') as fh:
            for line in fh:
                m = pat.match(line)
                if not m:
                    continue
                ip = m.group(1)
                path = m.group(2).split('?')[0]
                unique_ips.add(ip)
                total += 1
                counter[path] += 1
    except FileNotFoundError:
        return None, None, None, None
    except OSError:
        return None, None, None, None
    landing_hits = counter.get('/', 0)
    top_pages = counter.most_common(10)
    return unique_ips, total, top_pages, landing_hits


@bp.route('/api/admin/stats')
@superadmin_required
@limiter.limit("30 per minute")
def admin_stats():
    """Return today's site statistics: nginx log metrics + DB registrations."""
    from datetime import date as _date
    unique_ips, total_req, top_pages, landing_hits = _parse_nginx_today()
    log_available = unique_ips is not None

    db = SessionLocal()
    try:
        today_start = datetime.combine(_date.today(), datetime.min.time())
        registrations_today = db.query(func.count(User.id)).filter(
            User.created_at >= today_start
        ).scalar() or 0
    finally:
        db.close()

    return jsonify({
        'ok': True,
        'log_available': log_available,
        'unique_ips': len(unique_ips) if log_available else None,
        'total_requests': total_req if log_available else None,
        'landing_hits': landing_hits if log_available else None,
        'top_pages': top_pages if log_available else None,
        'registrations_today': registrations_today,
    })


# ── Health check ───────────────────────────────────────────────────────────────

@bp.route('/health')
@limiter.limit('60 per minute')
def health():
    """Public health-check endpoint. Returns JSON with uptime and DB status."""
    uptime_sec = int((datetime.utcnow() - _APP_START_TIME).total_seconds())
    db_status = "error"
    offers_count = 0
    try:
        db = SessionLocal.session_factory()
        try:
            offers_count = db.query(func.count(Offer.id)).scalar() or 0
            db_status = "ok"
        finally:
            db.close()
    except Exception:
        pass
    return jsonify({
        "status": "ok",
        "db": db_status,
    })


# ── Admin error log ────────────────────────────────────────────────────────────

def _get_log_path(key: str) -> Path:
    """Return the configured log path for 'errors' or 'security'."""
    from flask import current_app
    cfg_key = "ERROR_LOG_PATH" if key == "errors" else "SECURITY_LOG_PATH"
    return Path(current_app.config.get(cfg_key, str(DATA_DIR / f"{key}.log")))


def _read_error_log(n: int = 50) -> list[dict]:
    """Return the last *n* entries from errors.log (newest first)."""
    from collections import deque
    log_path = _get_log_path("errors")
    if not log_path.exists():
        return []
    try:
        lines = deque(log_path.read_text(encoding="utf-8").splitlines(), maxlen=n * 2)
        entries = []
        for line in reversed(lines):
            line = line.strip()
            if not line:
                continue
            try:
                entries.append(json.loads(line))
            except Exception:
                pass
            if len(entries) >= n:
                break
        return entries
    except Exception:
        return []


def _read_security_log(n: int = 200) -> list[str]:
    """Return the last *n* lines from security.log (newest first)."""
    from collections import deque
    log_path = _get_log_path("security")
    if not log_path.exists():
        return []
    try:
        lines = [l for l in log_path.read_text(encoding="utf-8").splitlines() if l.strip()]
        return list(reversed(lines[-n:]))
    except Exception:
        return []


@bp.route('/admin/errors')
@superadmin_required
@limiter.limit("30 per minute")
def admin_errors():
    """Admin error log: last 50 500 errors."""
    entries = _read_error_log(50)
    return render_template(
        'errors.html',
        entries=entries,
        role=get_current_role(),
        username=session.get("username"),
    )


@bp.route('/admin/errors/clear', methods=['POST'])
@superadmin_required
@limiter.limit("30 per minute")
def admin_errors_clear():
    """Truncate the errors log file."""
    try:
        _get_log_path("errors").write_text("", encoding="utf-8")
        return jsonify({"ok": True})
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)}), 500


@bp.route('/admin/security-log')
@superadmin_required
@limiter.limit("30 per minute")
def admin_security_log():
    """Admin security event log: last 200 lines."""
    lines = _read_security_log(200)
    return render_template(
        'security_log.html',
        lines=lines,
        role=get_current_role(),
        username=session.get("username"),
    )


@bp.route('/admin/security-log/clear', methods=['POST'])
@superadmin_required
@limiter.limit("30 per minute")
def admin_security_log_clear():
    """Truncate the security log file."""
    try:
        log_path = _get_log_path("security")
        log_path.write_text("", encoding="utf-8")
        # Re-add the file handler so new events can be written
        for h in _sec_logger.handlers:
            h.close()
        _sec_logger.handlers.clear()
        _init_security_logger(str(log_path))
        return jsonify({"ok": True})
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)}), 500


# ── Legal pages (public) ───────────────────────────────────────────────────────

@bp.route('/faq')
def faq():
    """FAQ & Guide de démarrage — public page."""
    return render_template('faq.html')


@bp.route('/cgu')
def cgu():
    """Conditions Générales d'Utilisation — public page."""
    return render_template('cgu.html')


@bp.route('/confidentialite')
def confidentialite():
    """Politique de Confidentialité (RGPD) — public page."""
    return render_template('confidentialite.html')


@bp.route('/mentions-legales')
def mentions_legales():
    """Mentions légales — public page."""
    return render_template('mentions_legales.html')


# ── Personal data export ───────────────────────────────────────────────────────

@bp.route('/api/account/export', methods=['POST'])
@login_required
@limiter.limit("5 per minute")
def account_export():
    """Export all personal data for the logged-in user as a JSON file (RGPD Article 20)."""
    import json as _json

    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'error': 'Non autorisé'}), 403

    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            return jsonify({'error': 'Utilisateur introuvable'}), 404

        # ── Profile data ─────────────────────────────────────
        export_data = {
            'profil': {
                'identifiant':   user.username,
                'email':         user.email,
                'role':          user.role,
                'domaine_id':    user.domain_id,
                'domaine':       user.domain.name if user.domain else None,
                'a2f_active':    user.totp_enabled,
                'inscription':   user.created_at.isoformat() if user.created_at else None,
                'derniere_connexion': user.last_login.isoformat() if user.last_login else None,
                'tokens_ia_utilises': user.claude_tokens_used,
                'nb_matchings':  user.matching_count,
            },
            'candidatures': [],
            'documents': [],
        }

        # ── UserOffer tracking data ───────────────────────────
        user_offers = (
            db.query(UserOffer)
            .filter(UserOffer.user_id == user_id)
            .options(joinedload(UserOffer.offer))
            .all()
        )
        for uo in user_offers:
            o = uo.offer
            export_data['candidatures'].append({
                'offre_id':     o.id if o else None,
                'titre':        o.title if o else None,
                'entreprise':   o.company if o else None,
                'url':          o.url if o else None,
                'statut':       uo.status,
                'cv_envoye':    uo.cv_sent,
                'relance_faite': uo.follow_up_done,
                'date_envoi':   uo.date_sent.isoformat() if uo.date_sent else None,
                'date_relance': uo.follow_up_date.isoformat() if uo.follow_up_date else None,
                'notes':        uo.notes,
                'score_match':  uo.cv_match_score,
            })

        # ── Documents list (filenames only, not binary content) ──
        docs_dir = DATA_DIR / 'documents' / str(user_id)
        if docs_dir.exists():
            export_data['documents'] = sorted(f.name for f in docs_dir.iterdir() if f.is_file())

        export_json = _json.dumps(export_data, ensure_ascii=False, indent=2, default=str)
        safe_username = "".join(c if c.isalnum() or c in "-_" else "_" for c in user.username)
        filename = f"myjobhunter_export_{safe_username}.json"

        from flask import Response
        return Response(
            export_json,
            mimetype='application/json',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Type': 'application/json; charset=utf-8',
            },
        )
    finally:
        db.close()
